from flask import Flask,render_template,jsonify,request, send_file
import pickle
import os
import joblib
import numpy as np 
from flask_cors import CORS
import numpy as np
from fpdf import FPDF
from docx import Document
import pandas as pd

# Application configuration:- 
app = Flask(__name__)

# Model configurations 
# Load the model
with open('model.pkl', 'rb') as model_file: 
    pipeline = pickle.load(model_file)
#pipeline = joblib.load("model.pkl")
@app.route('/',methods=["GET","POST"])
def index():
    return render_template("index.html")

@app.route('/predict',methods=["GET","POST"])
def predict():
    if request.method =='POST':
        # Get form data
        #data = request.form.to_dict()
        #Get JSON data from the request 
        #data = request.get_json(force=True)
        #data = [float(x) for x in request.form.values()]
        #return data
        # new_data = [float(x) for x in data]
        # Extract features from the JSON data 
        # features = [ int(data['person_age']), 
        #             int(data['person_gender']), 
        #             int(data['person_education']), 
        #             int(data['person_income']), 
        #             int(data['person_emp_exp']), 
        #             int(data['person_home_ownership']), 
        #             int(data['loan_amnt']), 
        #             int(data['loan_intent']), 
        #             float(data['loan_int_rate']), 
        #             int(data['loan_percent_income']), 
        #             int(data['cb_person_cred_hist_length']), 
        #             int(data['credit_score']), 
        #             int(data['previous_loan_defaults']) ]
        #return new_data
        # Make prediction
        # # Convert features to a numpy array 
        # features_array = np.array([data])
        # # Make prediction and get probabilities 
        # #prediction = model.predict(features_array) 
        # probabilities = model.predict_proba(features_array)
        # prediction_result = model.predict([features_array])
        data = [float(x) for x in request.form.values()]
        # Make prediction
        prediction_result = pipeline.predict([data])
        probabilities = pipeline.predict_proba([data])
        #return {'probability':probabilities}
        #return jsonify(prediction_result)
        loan_amnt = request.form.get('loan_amnt')
        loan_int_rate = request.form.get('loan_int_rate')
        # [[0.95, 0.05]]  # Meaning 95% probability for class 0 and 5% probability for class 1
        response = {'prediction': int(prediction_result[0]), 
                    'probability': float(probabilities[0][1]),  # Probability of default (assuming 1 means default)
                   'loan_amount': float(loan_amnt), 
                   'interest_rate': float(loan_int_rate), 
                   'loan_term': int(12) # Assuming loan_term is part of your form data 
                }
        return render_template('predict.html',response=response)
        #return render_template('index.html',prediction_result=prediction_result[0],probabilities=probabilities)
        #return render_template('result.html', prediction=prediction_result,probabilities=probabilities)
    return render_template('predict.html')
@app.route('/api/predict', methods=['GET','POST']) 
def predict_api(): # Get form data as a dictionary 
    if request.method == 'POST':
        # Get form data as a dictionary
        #data = request.get_json(force=True) # {"data":[23]}
        #data=request.json['data'] # {"data":{'person_age':23}}
        data = request.form.to_dict() # body form output format {"person_age":'23'}
        #data = request.form['person_age']
        #return jsonify(data)
        # Convert form data to features array 
        features = [int(data['person_age']),
                    int(data['person_gender']),
                    int(data['person_education']),
                    float(data['person_income']),
                    float(data['person_emp_exp']),
                    int(data['person_home_ownership']),
                    float(data['loan_amnt']),
                    int(data['loan_intent']),
                    float(data['loan_int_rate']),
                    float(data['loan_percent_income']),
                    float(data['cb_person_cred_hist_length']),
                    float(data['credit_score']),
                    int(data['previous_loan_defaults'])]
        
        #data['credit_score'] = int(data['credit_score'])
        
        # Convert features to a numpy array
        #features = np.array(data['data']).astype(np.float32)
        #features_array = np.array([features]).reshape(1,-1)
        #return jsonify([features])
        
        # Make prediction and get probabilities 
        features_array = list(features)
        #return jsonify(features_array)
        
        prediction = pipeline.predict([features_array]) 
        probabilities = pipeline.predict_proba([features_array])
        
        # Convert numpy arrays to lists for JSON serialization 
        prediction = prediction.tolist() 
        probabilities = probabilities.tolist()

        #return jsonify({'predictions': prediction, 'probabilities':probabilities}),201
        # Prepare the response 
        response = {'prediction': int(prediction[0]), 
                    'probability': float(probabilities[0][1]),  # Probability of default (assuming 1 means default)
                    'loan_amount': data['loan_amnt'], 
                    'interest_rate': float(data['loan_int_rate']), 
                    'loan_term': int(data['loan_term']) # Assuming loan_term is part of your form data 
                } 
        return jsonify(response)
    return render_template('prediction.html')

@app.route('/download-report/<filetype>', methods=['POST']) 
def download_report(filetype): 
    data = request.json 
    filename = f"report.{filetype}" 
    if filetype == 'csv': 
        df = pd.DataFrame([data]) 
        df.to_csv(filename, index=False) 
    elif filetype == 'xlsx': 
        df = pd.DataFrame([data]) 
        df.to_excel(filename, index=False) 
    elif filetype == 'pdf': 
        pdf = FPDF() 
        pdf.add_page() 
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 10, 'Loan Prediction Report', 0, 1, 'C') 
        pdf.ln(10)
        pdf.set_fill_color(200, 220, 255) 
        pdf.set_text_color(0) 
        pdf.set_draw_color(50, 50, 100) 
        pdf.set_line_width(0.3)
        for key, value in data.items(): 
            #pdf.cell(200, 10, txt=f"{key}: {value}", ln=True) 
            pdf.cell(0, 10, f"{key}: {value}", 0, 1, 'L', True)
        pdf.output(filename) 
    elif filetype == 'docx': 
        doc = Document() 
        doc.add_heading('Loan Prediction Report', 0) 
        table = doc.add_table(rows=1, cols=2) 
        hdr_cells = table.rows[0].cells 
        hdr_cells[0].text = 'Field' 
        hdr_cells[1].text = 'Value'
        for key, value in data.items(): 
            #doc.add_paragraph(f"{key}: {value}") 
            row_cells = table.add_row().cells 
            row_cells[0].text = str(key) 
            row_cells[1].text = str(value)
        doc.save(filename) 
    return send_file(filename, as_attachment=True)

if __name__=="__main__":
    app.run(debug=True,port=8080)
    



